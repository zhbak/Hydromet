from docxtpl import DocxTemplate
from docx import Document
from geopy import distance
from docx.shared import Inches
from dotenv import load_dotenv
from openai import OpenAI
from io import BytesIO
from PIL import Image
import pandas as pd
import requests, os, datetime

# Функция, которая будет обновлять и сохранять doc
# template  – шаблон файла на вход
# context - словарь
# report - имя или путь для нового обновленного файла 
def doc_report(template, context, report):
    doc = DocxTemplate(template)
    doc.render(context)
    doc.save(report)

# Cортировки по удалённости
# Возвращает dataframe с отсортированными станциями 
def distanse_sorting(object_latitude, object_longitude, stations_path):
    stations = pd.read_csv(stations_path, sep=';')
    
    # Функция для расчёта расстояния до метеостанции
    def distanse_from_object(row):
        return round(distance.geodesic((object_latitude, object_longitude), (row["Latitude"], row["Longitude"])).km, 1)
    
    # Функция для расчёта периода наблюдений
    def period(row):
        return datetime.datetime.now().year - int(row['ОТКРЫТИЕ'][-4:])

    stations["Distance_from_object"] = stations.apply(distanse_from_object, axis=1)
    stations["Period"] = stations.apply(period, axis=1)
    stations_100_km = stations[stations["Distance_from_object"] <= 100].sort_values(by='Distance_from_object')
    return stations_100_km
    
# Карта метостанций
def station_map(mapbox_access_token, save_path, latitude, longitude, map_url):

    base_url = f"{map_url}/{longitude},{latitude},7,0/700x700?access_token={mapbox_access_token}"

    # Выполнение запроса
    response = requests.get(base_url)

    # Конвертация изображения в формат PIL
    image = Image.open(BytesIO(response.content))
    image = image.convert('RGB')
    image.save(save_path, format='JPEG')

# Вставка таблицы и карты в .docx
def table_map_paste(object_latitude, object_longitude,
                    meteo_stations_path, meteo_station_map_path,
                    hydro_stations_path, hydro_station_map_path,
                    path_doc_input, path_doc_output):
    
    # Считывание данных из CSV и фильтрация до 100 км
    meteo_stations_100_km = distanse_sorting(object_latitude, object_longitude, meteo_stations_path)
    hydro_stations_100_km = distanse_sorting(object_latitude, object_longitude, hydro_stations_path)
    print("Distance calculated.")

    # Открытие документа
    doc = Document(path_doc_input)

    # Получение таблицы из документа
    table_meteo_0 = doc.tables[0]  
    table_meteo_1 = doc.tables[1]
    table_hydro_0 = doc.tables[2]
    table_hydro_1 = doc.tables[3]

    # Перебор данных из DataFrame (meteo) и заполнение таблицы в документе
    for index, row in meteo_stations_100_km.iterrows():
        new_row_0 = table_meteo_0.add_row().cells
        new_row_1 = table_meteo_1.add_row().cells
        new_row_0[0].text = str(row['НАЗВАНИЕ'])  
        new_row_0[1].text = str(row['ВМО'])
        new_row_0[2].text = str(row['ТИП'])  
        new_row_0[3].text = str(row['УГМС'])  
        new_row_0[4].text = str(row['Latitude'])  
        new_row_0[5].text = str(row['Longitude'])  
        new_row_0[6].text = str(row['ВЫС'])

        new_row_1[0].text = str(row['НАЗВАНИЕ'])  
        new_row_1[1].text = str(row['ВМО'])
        new_row_1[2].text = str(row['ОТКРЫТИЕ'])    
        new_row_1[3].text = str(row['ЗАКРЫТИЕ'])  
        new_row_1[4].text = str(row['Period'])   
        new_row_1[5].text = str(row['ИЗМЕНЕНИЯ']) 
        new_row_1[6].text = str(row['Distance_from_object'])

        # Перебор данных из DataFrame (hydro) и заполнение таблицы в документе
    for index, row in hydro_stations_100_km.iterrows():
        new_row_0 = table_hydro_0.add_row().cells
        new_row_1 = table_hydro_1.add_row().cells
        new_row_0[0].text = str(row['Наименование поста'])  
        new_row_0[1].text = str(row['Код поста'])
        new_row_0[2].text = str(row['Расстояние от устья, км'])  
        new_row_0[3].text = str(row['Площадь водосбора, кв.км.'])  
        new_row_0[4].text = str(row['Latitude'])  
        new_row_0[5].text = str(row['Longitude'])  
        new_row_0[6].text = str(row['Высота, м'])
        new_row_0[7].text = str(row['Система высот'])

        new_row_1[0].text = str(row['Наименование поста'])  
        new_row_1[1].text = str(row['Код поста'])
        new_row_1[2].text = str(row['ОТКРЫТИЕ'])    
        new_row_1[3].text = str(row['Закрыт'])  
        new_row_1[4].text = str(row['Period'])   
        new_row_1[5].text = str(row['Принадлежность поста']) 
        new_row_1[6].text = str(row['Distance_from_object'])    

    # Вставка карты метеостанций
    doc.add_picture(meteo_station_map_path, width=Inches(5.5))
    doc.add_paragraph("Рисунок – Метеорологическая изученность")
    doc.add_picture(hydro_station_map_path, width=Inches(5.5))
    doc.add_paragraph("Рисунок – Гидрологическая изученность")

    # Подгружаем ключи
    load_dotenv()
    client = OpenAI(
                    # This is the default and can be omitted
                    api_key=os.environ.get("OPENAI_API_KEY")
                    )
    
    with open("hydromet_stations/SP_information.txt", encoding="utf8") as file:
        SP_information = file.read()
        file.close()

    request = "{} Таблица метеорологической изученности:\n{}\nТаблица гидрологической изученности:{}".format(SP_information, 
                                                                                                             meteo_stations_100_km.to_string(), 
                                                                                                             hydro_stations_100_km.to_string()) 
    message = [
                {
                 "role" : "user",
                 "content" : request   
                }
                ]
    response = client.chat.completions.create(
                                      model='gpt-3.5-turbo-16k',
                                      messages=message,
                                      temperature=0.4
                                    )
    
    doc.add_paragraph(response.choices[0].message.content)

    # Сохранение обновленного документа
    doc.save(path_doc_output)  