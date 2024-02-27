from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
import os, convertapi

# Подгружаем ключи
load_dotenv()
client = OpenAI(
    # This is the default and can be omitted
    api_key=os.environ.get("OPENAI_API_KEY"),
)

KEY_convert = os.environ.get("API_Secret_convert")

# Функция для конвертации .pdf в .txt через API convert
def pdfocr_to_txt(KEY_convert, input, output):
    convertapi.api_secret = KEY_convert
    convertapi.convert('txt', {
        'File': input,
        'OcrLanguage' : 'ru'
        }, from_format = 'pdf').save_files(output)

# Функция читающая файл .docx
def read_docx(file_path):
    print("TR.docx is reading.")
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    print("TR.docx was read.")
    return '/n'.join(full_text)

# Функция сравнения docx и pdf
def file_compare(file_path):
    # Проверяем формат ТЗ (.pdf или .docx)
    print("File format is detectioning.")
    if '.docx' in file_path:
        print("File format is. docx.")
        #Применям функцию трансформации для .docx
        text_for_chat_gpt = read_docx(file_path)
        print("Text from TR.docx for chatgpt message was made.")

    elif '.pdf' in file_path:
        print("File format is .pdf.")
        # Применям функцию трансформации для .pdf
        print("Transformation from TR.pdf to TR.txt.")
        pdfocr_to_txt(KEY_convert, file_path, 'introduction/output')
        print("TR.txt is reading.")
        with open (os.path.join('introduction/output', 'TR.txt'), 'r', encoding='utf-8', errors='ignore') as file:
            text_for_chat_gpt = file.read().strip()
        print("Text from TR.pdf for chatgpt message was made.")   
    
    return text_for_chat_gpt

# Функция задающая prompt
def TR_introduction(TR_content):
    print("Asking ChatGPT.")
    TR_content_res = "ТЗ :{} \n".format(TR_content) + \
                     '''
                        инженерно-гидрометеорологические/их изыскания/ий = ИГМИ
                        
                        На основе ТЗ напиши введение для отчёта по инженерно-гидрометеорологическим
                        изысканиям, которое должно включать следующие параметры:
                            Наименование_объекта
                            Местоположение_объекта
                            Цели
                            Задачи
                            Сроки_выполнения_ИГМИ
                            Основание_для_выполнения_ИГМИ 
                            Вид_градостроительной_деятельности
                            Этап_выполнения_ИГМИ
                            Идентификационные_сведения_об_объекте
                            Характеристика_объекта
                            Сведения_о_заказчике
                            Сведения_об_исполнителе_работ
                            Сведения_о_ранее_выполненных_ИГМИ
                            ИГМИ_соответствуют_требованиям

                        Эти параметры в ТЗ часто представлены в неявном виде. Ты должен выявить их значения. Значения должны быть развернутым тестом.
                        Ответ приведи в виде python словаря {"параметр":"значение"}. При этом "значение" должно быть текстом, а не словарём.
                        Наименование может быть длинным и включать кадастровые номера и адрес.
                        Идентификационные сведения об объекте, как правило, включают уровень ответственности и класс, приводи также остальные характеристикии.
                        Характеристика объекта обычна включает его площадь, габариты и другие параметры, если её нет, то пиши Отсутствует.
                        Исполнитель работ обычно указан в блоке Согласовано, также исполнитель может быть подрядчиком.
                        В значениях исполнитель и заказчик указывай также и остальную информацию о них, включая адрес, фамилию и инициалы директора и т.д.
                        Сроки_выполнения_ИГМИ - если нет в ТЗ, то пиши Согласно договору.
                        ИГМИ_соответствуют_требованиям - нужно перечислить все нормативные документы включая их названия, если их нет, то пиши 'Согласно СП'.
                     '''
    message = [
                {
                 "role" : "user",
                 "content" : TR_content_res   
                }
                ]
    response = client.chat.completions.create(
                                      model='gpt-3.5-turbo-16k',
                                      messages=message,
                                      temperature=0.2
                                    )
    print("ChatGPT answered.")
    return(response.choices[0].message.content)


'''
# Старые функции
# Функция трансформирующая pdf в image
def pdf_to_image(path, path_pdf):
    
    # Создаем папку для хранения фото
    output_folder = os.path.join(path, "output_images/")
    shutil.rmtree(output_folder)
    os.mkdir(output_folder)

    # Конвертация
    images = convert_from_path(path_pdf, dpi=900)

    # Сохранение jpg
    for i, image in enumerate(images):
        images_path = os.path.join(output_folder, f"{i}.jpg")
        image.save(images_path, "JPEG")

# Функция распознования текста
def text_recogntion(path_with_images):

    # Получаем список изображений
    image_files = [file for file in os.listdir(path_with_images) if file.endswith(".jpg")]
    print(image_files)

    # Создаем экземпляр EasyOCR
    render = easyocr.Reader(['ru'])

    recognized_texts = ""

    # Циклом по файлам изображений
    for image_file in image_files:
        image_path = os.path.join(path_with_images, image_file)
        print(image_path)
        
        # Распознование текста на изображении
        recognized_text = render.readtext(image_path, detail=0, paragraph = True)
        recognized_text_union = ' '.join(recognized_text)
        recognized_texts += recognized_text_union
    
    return recognized_texts     

'''
